"""Render meeting minutes to a Word document via python-docx (EXP-1).

User customisation
------------------
Drop a ``.docx`` file at ``templates/export/docx_template.docx`` to
inherit its paragraph / heading styles. python-docx opens the template
(if present) and reuses its style table; otherwise it falls back to the
library defaults.

The exporter consumes the stored markdown minutes plus the optional
transcript. Markdown parsing here is deliberately minimal — we only map
headings, bullet/numbered lists, and action items rendered as
checkboxes. Anything more exotic gets flattened to a paragraph.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import TYPE_CHECKING

from meeting_minutes.export import ExportDependencyMissing
from meeting_minutes.system3.db import ActionItemORM, MeetingORM

if TYPE_CHECKING:  # pragma: no cover
    from docx.document import Document  # type: ignore


def _require_docx():
    try:
        import docx  # type: ignore
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise ExportDependencyMissing(
            "Install python-docx to enable DOCX export: pip install python-docx"
        ) from exc
    return docx, Document, Pt


def _template_path() -> Path:
    return (
        Path(__file__).resolve().parent.parent.parent.parent
        / "templates"
        / "export"
        / "docx_template.docx"
    )


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_ULIST_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_OLIST_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_CHECKBOX_RE = re.compile(r"^\s*[-*+]\s+\[([ xX])\]\s+(.*)$")


def _add_paragraph(doc, text: str) -> None:
    doc.add_paragraph(text)


def _add_bullets(doc, items: list[str], *, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        try:
            doc.add_paragraph(item, style=style)
        except KeyError:
            # Template missing the list style — fall back to plain paragraph.
            doc.add_paragraph(("1. " if ordered else "• ") + item)


def _flush_lists(doc, ul: list[str], ol: list[str]) -> tuple[list[str], list[str]]:
    if ul:
        _add_bullets(doc, ul, ordered=False)
    if ol:
        _add_bullets(doc, ol, ordered=True)
    return [], []


def _render_markdown_body(doc, markdown_text: str) -> None:
    """Walk the stored markdown a line at a time, mapping structure to docx.

    Skips the top-level ``# Title`` because the caller already wrote a
    Title paragraph. Everything else becomes Heading 2/3/…, bulleted list,
    numbered list, or paragraph.
    """
    lines = markdown_text.splitlines()
    ul_buf: list[str] = []
    ol_buf: list[str] = []
    skipped_first_h1 = False

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            ul_buf, ol_buf = _flush_lists(doc, ul_buf, ol_buf)
            continue

        m = _HEADING_RE.match(line)
        if m:
            ul_buf, ol_buf = _flush_lists(doc, ul_buf, ol_buf)
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
                continue
            heading_level = min(level, 4)
            doc.add_heading(text, level=heading_level)
            continue

        m = _CHECKBOX_RE.match(line)
        if m:
            ol_buf = _flush_lists(doc, [], ol_buf)[1]
            checked, text = m.group(1), m.group(2).strip()
            marker = "☑" if checked.lower() == "x" else "☐"
            ul_buf.append(f"{marker}  {text}")
            continue

        m = _ULIST_RE.match(line)
        if m:
            ol_buf = _flush_lists(doc, [], ol_buf)[1]
            ul_buf.append(m.group(1).strip())
            continue

        m = _OLIST_RE.match(line)
        if m:
            ul_buf = _flush_lists(doc, ul_buf, [])[0]
            ol_buf.append(m.group(1).strip())
            continue

        # Plain paragraph.
        ul_buf, ol_buf = _flush_lists(doc, ul_buf, ol_buf)
        _add_paragraph(doc, line.strip())

    _flush_lists(doc, ul_buf, ol_buf)


def _add_action_items_table(doc, actions: list[ActionItemORM]) -> None:
    """Render ActionItemORM rows as a 5-column Word table."""
    if not actions:
        return
    doc.add_heading("Action Items", level=2)
    table = doc.add_table(rows=1, cols=5)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:  # pragma: no cover - template lacks the style
        pass
    hdr = table.rows[0].cells
    hdr[0].text = "Description"
    hdr[1].text = "Owner"
    hdr[2].text = "Due"
    hdr[3].text = "Priority"
    hdr[4].text = "Status"
    for ai in actions:
        row = table.add_row().cells
        row[0].text = ai.description or ""
        row[1].text = ai.owner or ""
        row[2].text = ai.due_date or ""
        row[3].text = ai.priority or ""
        row[4].text = ai.status or ""


def render_docx(meeting: MeetingORM, *, with_transcript: bool = False) -> bytes:
    """Compose a ``.docx`` byte stream for ``meeting``."""
    _, Document, Pt = _require_docx()

    tpl = _template_path()
    doc = Document(str(tpl)) if tpl.exists() else Document()

    # Title + metadata.
    title = (meeting.title or "Meeting").strip()
    doc.add_heading(title, level=0)

    date_str = meeting.date.strftime("%Y-%m-%d") if meeting.date else ""
    duration = meeting.duration or ""
    attendees = ", ".join(a.name for a in (meeting.attendees or []) if a and a.name)

    meta_parts = []
    if date_str:
        meta_parts.append(f"Date: {date_str}")
    if duration:
        meta_parts.append(f"Duration: {duration}")
    if attendees:
        meta_parts.append(f"Attendees: {attendees}")
    if meta_parts:
        p = doc.add_paragraph()
        run = p.add_run("   ·   ".join(meta_parts))
        run.italic = True

    # Body from the stored markdown — minus the duplicated top-level title.
    _render_markdown_body(doc, (meeting.minutes.markdown_content or "").strip())

    # Replace whatever free-form "Action Items" markdown had with a proper table.
    if meeting.action_items:
        _add_action_items_table(doc, list(meeting.action_items))

    if with_transcript and meeting.transcript is not None and meeting.transcript.full_text:
        doc.add_page_break()
        doc.add_heading("Full Transcript", level=1)
        # python-docx doesn't ship a monospace style out of the box; fall back
        # to a plain paragraph split on blank lines.
        for para in meeting.transcript.full_text.strip().split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
